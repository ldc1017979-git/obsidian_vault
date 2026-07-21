from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).parent
DOCX = OUT / "库位占用状态与现场拍照优化PRD-V0.9-评审稿.docx"

BLUE = "2458F5"; DARK_BLUE = "1F4D78"; NAVY = "183B66"; TEXT = "303846"
MUTED = "667085"; LIGHT = "F2F4F7"; PALE_BLUE = "EEF3FF"; BORDER = "D9DEE7"
RED = "B42318"; GOLD = "9A6700"; GREEN = "067647"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(TEXT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Title", 24, NAVY, 0, 8), ("Subtitle", 13, MUTED, 0, 14),
    ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK_BLUE, 8, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"; st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    if name.startswith("Heading"): st.font.bold = True

for list_name in ("List Bullet", "List Number"):
    st = styles[list_name]
    st.font.name = "Calibri"; st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.paragraph_format.left_indent = Inches(0.5); st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(4); st.paragraph_format.line_spacing = 1.167

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None: tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[idx])); tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cell)

def add_table(headers, rows, widths, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.text=str(h); set_cell_shading(cell,header_fill)
        for r in cell.paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    for row in rows:
        cells=table.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
    set_table_geometry(table,widths)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0)
    return table

def add_bullets(items):
    for item in items: doc.add_paragraph(item, style="List Bullet")

def add_numbered(items):
    for item in items: doc.add_paragraph(item, style="List Number")

def add_callout(label, text, fill=PALE_BLUE, color=NAVY, trailing=True):
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; set_table_geometry(t,[9360])
    c=t.cell(0,0); set_cell_shading(c,fill); c.text=""
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label+"："); r.bold=True; r.font.color.rgb=RGBColor.from_string(color)
    p.add_run(text)
    if trailing:
        doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold_prefix=None):
    p=doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); r.bold=True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    return p

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement("w:tblHeader"); tblHeader.set(qn("w:val"), "true"); trPr.append(tblHeader)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    cant.set(qn("w:val"), "true")
    trPr.append(cant)

# Header/footer
hp=sec.header.paragraphs[0]; hp.text="产品需求文档（PRD）｜库位占用状态与现场拍照优化"; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
for r in hp.runs: r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(MUTED)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
fp.add_run("内部评审稿  |  ")
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); fp._p.append(fld)
for r in fp.runs: r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(MUTED)

# Masthead
p=doc.add_paragraph(style="Title"); p.add_run("库位占用状态与现场拍照优化")
p=doc.add_paragraph(style="Subtitle"); p.add_run("平板端叉车员批量修改库位状态 × PC端最新占用凭证展示")
add_table(["文档属性","内容"],[
    ("版本","V0.9（评审稿）"),("日期","2026-07-21"),("状态","待产品、研发、测试联合评审"),
    ("产品范围","叉车员平板端、库位PC端、照片与操作记录"),("优先级","P1"),("作者","Codex整理，产品负责人待确认")
],[2400,6960],header_fill=PALE_BLUE)
add_callout("评审声明","本文件根据本轮已确认需求、PC端原型、平板端原型及端到端流程图汇总形成。当前为评审稿；通过人工评审和PRD门禁后方可标记正式版本。")

add_heading("1. 需求背景与目标",1)
add_heading("1.1 背景",2)
add_para("现有叉车员库位操作缺少强制现场照片凭证，管理人员在PC端也无法直接查看本次占用操作的人员、时间与现场图片，导致库存差异核查、责任追溯和异常复盘效率较低。")
add_para("本次优化同时覆盖平板端与PC端：叉车员在平板端批量选择空闲库位，将其统一修改为占用状态并拍摄现场照片；管理人员在PC端仅针对占用状态查看最新一次操作凭证。")
add_heading("1.2 目标",2)
add_bullets([
    "保证空闲库位转占用时至少留存1张现场照片，形成可追溯凭证。",
    "支持叉车员一次选择多个空闲库位并批量修改，减少重复操作。",
    "严格限制状态权限：叉车员只能编辑空闲库位，且只能改为占用。",
    "PC端仅在占用状态展示最新叉车员、上传时间和照片，避免历史信息干扰当前判断。",
    "后台对单个业务对象滚动保留最近3次记录，兼顾追溯需求与存储控制。"
])
add_heading("1.3 成功标准",2)
add_table(["指标","目标","验证方式"],[
    ("现场凭证完整率","所有成功的空闲→占用操作均有1～3张照片","抽查操作记录与图片关联"),
    ("非法状态修改率","占用、锁闭等非空闲状态修改成功数为0","权限与状态用例"),
    ("批量操作一致性","一次提交中的全部库位要么全部成功，要么全部失败","并发/事务测试"),
    ("PC展示准确率","占用状态展示最新记录；其他状态不展示","跨端验收"),
],[2600,2600,4160])

add_heading("2. 范围定义",1)
add_heading("2.1 本期范围",2)
add_bullets([
    "平板端叉车员登录、车间/库区选择、库位图查看与空闲库位多选。",
    "“修改库位状态”入口及批量修改弹窗。",
    "目标状态固定为占用；现场拍照1～3张；照片预览和删除。",
    "提交校验、状态冲突处理、成功反馈与库位图刷新。",
    "PC端占用状态下展示最新叉车员、上传时间和照片。",
    "每个业务对象最多保留最近3次有效上传记录。"
])
add_heading("2.2 明确不在本期",2)
add_bullets([
    "叉车员将占用、锁闭或其他状态修改为空闲或其他状态。",
    "平板端展示库位详情、历史叉车员、历史上传时间或历史照片。",
    "从相册或本地文件选择照片。",
    "PC端直接展示最近3次历史记录；PC端只显示最新一次。",
    "后台历史记录查询页面、照片人工审核、图片标注和OCR识别。"
])

add_heading("3. 用户、角色与权限",1)
add_table(["角色","端","可执行动作","限制"],[
    ("叉车员","平板端","登录、切换车间/库区、多选空闲库位、现场拍照、删除照片、批量提交","不可编辑非空闲库位；不可选择目标状态；不可从相册选图"),
    ("仓储管理人员","PC端","查看库位当前状态；占用时查看最新叉车员、时间与照片","仅展示权限范围内库位；不展示非占用状态凭证"),
    ("系统","服务端","校验、批量更新、保存记录、滚动清理、返回最新凭证","不得部分成功；不得在新记录失败时删除旧记录"),
],[1600,1200,3480,3080])

add_heading("4. 业务对象与状态",1)
add_table(["业务对象","职责","关键标识","关系"],[
    ("库位","承载当前可用状态","location_id","属于车间/库区；关联多次操作记录"),
    ("批量状态修改","承载一次多库位操作","operation_id / request_id","包含1个或多个库位，共用1～3张照片"),
    ("上传记录","保存操作人、时间与凭证","evidence_id","与批量操作及库位关联"),
    ("照片","现场凭证","file_id","一次操作1～3张，按顺序展示"),
],[1800,3000,2200,2360])
add_heading("4.1 库位状态权限矩阵",2)
add_table(["当前状态","平板端可选","可修改","允许目标状态","PC占用字段"],[
    ("空闲","是","是","占用","不展示"),
    ("占用","否","否","不适用","展示最新叉车员、上传时间、照片"),
    ("锁闭","否","否","不适用","不展示"),
    ("其他状态","否","否","不适用","不展示"),
],[1500,1400,1400,2100,2960])

add_heading("5. 核心业务规则",1)
rules=[
("R01","库位可编辑条件","仅当前状态为空闲的库位可被叉车员选择和编辑。"),
("R02","批量选择","一次操作至少选择1个空闲库位，可选择多个；选中库位支持再次点击取消。"),
("R03","目标状态","叉车员只能将所选空闲库位修改为占用，目标状态固定且不可切换。"),
("R04","入口状态","未选择任何库位时，“修改库位状态”按钮禁用。"),
("R05","拍照来源","仅允许调用设备相机现场拍照，不提供相册或文件选择入口。"),
("R06","照片数量","一次批量操作共用照片，最少1张、最多3张。"),
("R07","照片删除","提交前可删除照片；删除后剩余0张时不可提交。"),
("R08","无照片拦截","未上传照片点击完成时提示“请至少拍摄1张现场照片，未上传照片不可提交”，且保留弹窗。"),
("R09","上限拦截","第4次拍照不新增照片，提示“最多只能拍摄3张照片”。"),
("R10","批量原子性","一次提交中的全部库位统一成功或统一失败，不允许部分库位成功。"),
("R11","并发校验","提交时若任一库位已不再为空闲，则整批拒绝并提示刷新后重试。"),
("R12","PC展示条件","PC端仅当前状态为占用时展示叉车员、上传时间和照片。"),
("R13","PC最新记录","PC端只展示最新一次有效上传记录及其全部1～3张照片。"),
("R14","记录保留","每个业务对象最多保留最近3次有效记录；第4次成功保存后删除最早记录。"),
("R15","安全删除顺序","新记录保存成功后才允许清理最早记录；清理失败须记录并重试。"),
("R16","重复提交","同一request_id不得重复生成操作记录或重复更新库位。"),
]
add_table(["编号","规则","业务要求"],rules,[900,2200,6260],header_fill=PALE_BLUE)

add_heading("6. 平板端产品需求",1)
add_heading("6.1 登录",2)
add_bullets(["叉车员输入用户名和密码登录；登录后展示本人姓名。","登录失败提示原因，不进入库位操作页面。","退出后返回登录页，不保留未提交操作。"])
add_heading("6.2 库位操作主页面",2)
add_bullets([
    "顶部展示车间选择、当前库区、已选库位数量和“修改库位状态”按钮。",
    "中部展示库区页签及库位图；用明确颜色区分空闲、占用、锁闭和已选。",
    "页面不展示单个库位的叉车员、上传时间、照片或库存详情。",
    "点击空闲库位进行多选/取消；点击占用、锁闭或其他状态时提示“仅空闲状态库位可以修改”。",
    "切换车间或库区前，如存在已选库位，应提示确认并清空当前选择。"
])
add_heading("6.3 修改库位状态弹窗",2)
add_table(["区域","内容","规则"],[
    ("已选库位","数量及库位编号列表","只读；来源于主页面选择"),
    ("目标状态","占用","只读；不可选择其他状态"),
    ("现场拍照","相机入口、照片缩略图、删除","一次操作1～3张；无相册入口"),
    ("操作区","取消、完成并修改库位状态","无照片时点击完成必须提示并拦截"),
],[1800,3000,4560])
add_heading("6.4 提交与反馈",2)
add_numbered([
    "叉车员点击“完成并修改库位状态”。",
    "前端校验已选库位数量和照片数量。",
    "服务端重新校验角色权限、库位最新状态和照片有效性。",
    "全部校验通过后，批量将库位改为占用并保存操作记录。",
    "成功后关闭弹窗、清空选择、刷新库位图并提示修改数量。",
    "失败时保留选择和照片，展示可恢复的错误提示。"
])

add_heading("7. PC端产品需求",1)
add_heading("7.1 展示位置与条件",2)
add_para("在现有库位PC页面右侧统计区域下方增加占用凭证信息。仅当库位当前状态为占用时，展示叉车员、上传时间和照片；其他状态不展示字段名称、字段内容及照片区域。")
add_heading("7.2 字段定义",2)
add_table(["字段","来源","展示规则","格式/交互"],[
    ("叉车员","最新有效记录","占用状态展示","姓名；可附工号/账号"),
    ("上传时间","最新有效记录提交成功时间","占用状态展示","YYYY-MM-DD HH:mm:ss"),
    ("照片","最新有效记录","展示全部1～3张","缩略图；点击查看大图；支持前后切换"),
],[1500,2300,2700,2860])
add_heading("7.3 历史兼容",2)
add_para("上线前或异常迁移产生的占用记录若无照片，PC端显示“历史数据暂无照片”，页面必须可正常打开。该兼容仅适用于存量数据，新提交记录必须满足照片必填。")

add_heading("8. 端到端流程",1)
add_numbered([
    "叉车员登录平板端并选择车间、库区。",
    "叉车员多选一个或多个空闲库位；非空闲状态被拦截。",
    "叉车员点击“修改库位状态”，系统展示已选库位及固定目标状态“占用”。",
    "叉车员现场拍摄1～3张照片，可删除重拍。",
    "叉车员点击完成；无照片则拦截，有照片则进入服务端校验。",
    "服务端校验通过后批量更新状态、保存操作与照片记录。",
    "如记录超过3次，新记录成功后清理最早记录。",
    "PC端读取当前状态；占用则展示最新凭证，其他状态隐藏凭证字段。"
])
add_callout("流程图","配套可交互流程图：flowcharts/2026-07-21-库位占用状态修改端到端流程-flowchart.html")

add_heading("9. 数据与记录要求",1)
add_heading("9.1 产品级记录字段",2)
add_table(["对象","字段","必填","说明"],[
    ("批量操作","operation_id / request_id","是","唯一操作与幂等标识"),
    ("批量操作","operator_id / operator_name_snapshot","是","叉车员标识及姓名快照"),
    ("批量操作","submitted_at","是","成功提交时间"),
    ("库位关联","location_id","是","一次操作关联多个库位"),
    ("库位关联","before_status / after_status","是","空闲→占用"),
    ("照片","file_id / storage_key","是","照片文件标识或存储键"),
    ("照片","sort_order","是","1～3"),
],[1600,3000,1000,3760])
add_heading("9.2 最近3次记录",2)
add_bullets([
    "保留范围按业务对象计算，具体以location_id还是占用周期ID为准，见待研发确认项。",
    "记录按成功提交时间倒序；只保留最新3次有效记录。",
    "第4次成功后删除第1次，第5次成功后删除第2次，以此类推。",
    "数据库记录删除与照片实体文件清理应解耦，文件清理失败可重试且不得影响新记录。"
])

add_heading("10. 异常与边界",1)
add_table(["编号","场景","预期结果"],[
    ("E01","未选择库位","修改按钮禁用"),
    ("E02","选择占用/锁闭/其他状态","不选中，提示仅空闲状态可修改"),
    ("E03","未拍照点击完成","提示至少1张，弹窗保持打开"),
    ("E04","拍摄第4张","不新增，提示最多3张"),
    ("E05","删除至0张","允许继续拍照；提交时拦截"),
    ("E06","相机权限未开启","提示开启权限并提供去设置入口"),
    ("E07","照片上传失败","标记失败照片；成功照片无需重拍；不可提交"),
    ("E08","网络中断","保留选择和照片，允许重试"),
    ("E09","并发导致库位状态变化","整批失败，提示刷新后重新选择"),
    ("E10","重复点击提交","仅处理一次，不生成重复记录"),
    ("E11","新记录保存失败","不更新库位、不删除旧记录"),
    ("E12","历史照片加载失败","PC端显示加载失败占位，不影响其他字段"),
],[900,3300,5160])

add_heading("11. 非功能与安全要求",1)
add_bullets([
    "权限：服务端必须校验叉车员角色、数据范围和库位操作权限，不能仅依赖前端按钮隐藏。",
    "一致性：批量状态更新、操作记录和照片关联应具备一致性保障，不允许部分成功。",
    "幂等：客户端每次提交生成唯一request_id，重复请求返回同一处理结果。",
    "安全：照片不可使用永久公开地址；PC查看时校验权限或使用短时签名地址。",
    "性能：库位图加载、选择和弹窗打开应满足现场连续操作，具体指标待研发基线确认。",
    "审计：记录操作者、时间、库位集合、状态变化、照片和异常结果。"
])

add_heading("12. 验收标准（AC）",1)
acs=[
("AC01","空闲库位可多选，已选数量实时更新。"),("AC02","再次点击已选空闲库位可取消。"),
("AC03","占用、锁闭及其他状态不可选、不可修改。"),("AC04","未选库位时修改按钮禁用。"),
("AC05","目标状态固定为占用。"),("AC06","平板主页面不展示库位详情。"),
("AC07","修改弹窗展示全部已选库位。"),("AC08","弹窗仅提供现场拍照入口，无相册入口。"),
("AC09","拍摄1～3张时可正常预览。"),("AC10","第4张被阻止并提示上限。"),
("AC11","照片支持删除，数量同步更新。"),("AC12","未拍照点击完成时明确提示并阻止提交。"),
("AC13","删除至0张后点击完成仍被阻止。"),("AC14","合法提交后全部所选库位变为占用。"),
("AC15","任一库位状态冲突时整批失败。"),("AC16","成功后清空选择并刷新库位图。"),
("AC17","重复提交不产生重复更新或重复记录。"),("AC18","PC端占用状态展示最新叉车员。"),
("AC19","PC端占用状态展示最新上传时间。"),("AC20","PC端占用状态展示最新1～3张照片。"),
("AC21","PC端其他状态不展示上述字段。"),("AC22","PC照片支持大图及前后切换。"),
("AC23","每个业务对象最多保留最近3次记录。"),("AC24","第4次保存失败时原3次记录不被删除。"),
]
add_table(["编号","验收标准"],acs,[1000,8360],header_fill=PALE_BLUE)

add_heading("13. 需求追溯矩阵",1)
trace=[
("R01-R04","平板库位图/操作栏","AC01-AC07","空闲/占用/锁闭选择及入口状态"),
("R05-R09","修改状态弹窗","AC08-AC13","照片来源、数量、删除与必填"),
("R10-R11,R16","服务端提交","AC14-AC17,AC24","批量一致性、冲突、幂等"),
("R12-R13","PC占用凭证区","AC18-AC22","状态条件、最新记录、照片预览"),
("R14-R15","记录清理","AC23-AC24","最近3次与安全清理顺序"),
]
add_table(["规则","页面/模块","验收标准","测试重点"],trace,[1800,2200,2200,3160])

add_heading("14. 待确认与风险",1)
add_table(["编号","类型","待确认项","建议口径","阻塞性"],[
    ("P01","业务","最近3次按location_id还是占用周期ID计算","按单次占用周期ID；PC仍只展示最新","不阻塞原型，阻塞数据设计"),
    ("P02","业务","一次多选库位的照片是否同时关联所有库位","同一批次照片关联全部所选库位","阻塞数据设计"),
    ("P03","研发","批量提交的一致性与并发控制方案","保证整批成功或整批失败","阻塞开发"),
    ("P04","研发","单张图片格式、大小与压缩规则","JPG/JPEG/PNG，压缩后≤5MB","阻塞接口定稿"),
    ("P05","研发","相机来源真实性校验能力","前端只开放相机；服务端校验上传凭证","不阻塞原型"),
    ("P06","运维","被淘汰记录对应照片实体的清理与重试","异步清理、失败重试并记录日志","不阻塞前端"),
],[800,1100,3300,3160,1000])

add_heading("15. 评审与版本",1)
add_heading("15.1 评审角色",2)
add_table(["角色","评审重点","当前状态"],[
    ("产品负责人","范围、规则、优先级、待确认项","待评审"),
    ("业务代表","现场操作路径、状态权限、术语","待评审"),
    ("研发代表","接口、并发、一致性、记录清理","待评审"),
    ("测试代表","AC、异常、边界、跨端一致性","待评审"),
    ("设计/交互","平板与PC页面结构、操作反馈","已有原型，待签收"),
],[1800,5200,2360])
add_heading("15.2 版本记录",2)
add_table(["版本","日期","内容","状态"],[
    ("V0.1","2026-07-21","新增单库位拍照与PC凭证展示","历史草案"),
    ("V0.5","2026-07-21","调整为现场拍照、PC仅展示最新记录","历史草案"),
    ("V0.8","2026-07-21","平板端增加空闲库位多选、固定转占用","原型确认中"),
    ("V0.9","2026-07-21","汇总跨端、数据、异常与AC，形成完整评审稿","当前版本"),
],[1200,1700,4660,1800])
add_callout("门禁结论","资料和业务主线已完整，原型及交互测试已覆盖；但人工产品/业务/研发/测试评审尚未完成，因此当前结论为 REVIEW，不得标记正式PRD。",fill="FFF6E8",color=GOLD,trailing=False)

terminal_p = doc.add_paragraph()
terminal_p.paragraph_format.space_before = Pt(0)
terminal_p.paragraph_format.space_after = Pt(0)
terminal_p.paragraph_format.line_spacing = Pt(1)
terminal_p.add_run("").font.size = Pt(1)

doc.core_properties.title = "库位占用状态与现场拍照优化PRD"
doc.core_properties.subject = "叉车员平板端批量修改库位状态与PC端最新凭证展示"
doc.core_properties.author = "Codex（整理）"
doc.core_properties.comments = "V0.9评审稿"

for table in doc.tables:
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name="Calibri"; run._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei")
                    run.font.size=Pt(9)

doc.save(DOCX)
print(DOCX)
