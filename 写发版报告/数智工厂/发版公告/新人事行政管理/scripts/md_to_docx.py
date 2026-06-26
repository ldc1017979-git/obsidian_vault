# -*- coding: utf-8 -*-
"""将 Markdown 转为 Word（.docx），支持标题、表格、图片、代码块。"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_document_font(document: Document, font_name: str = "微软雅黑") -> None:
    """设置文档默认中文字体。"""
    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def parse_table_row(line: str) -> list[str]:
    """解析 Markdown 表格行。"""
    line = line.strip()
    if not line.startswith("|"):
        return []
    cells = [cell.strip() for cell in line.strip("|").split("|")]
    return cells


def is_table_separator(line: str) -> bool:
    """判断是否为表格分隔行。"""
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    """将 Markdown 表格写入 Word。"""
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_index, row_cells in enumerate(rows):
        for col_index in range(col_count):
            cell_text = row_cells[col_index] if col_index < len(row_cells) else ""
            table.rows[row_index].cells[col_index].text = cell_text


def add_rich_text(paragraph, text: str) -> None:
    """向段落写入含 **粗体** 的文本。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def resolve_image_path(md_dir: Path, image_path: str) -> Path | None:
    """解析相对图片路径。"""
    candidate = (md_dir / image_path).resolve()
    if candidate.exists():
        return candidate
    return None


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    """主转换逻辑。"""
    md_text = md_path.read_text(encoding="utf-8")
    md_dir = md_path.parent
    lines = md_text.splitlines()

    document = Document()
    set_document_font(document)

    index = 0
    in_code_block = False
    code_language = ""
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_language = stripped[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                title = f"[{code_language or '代码'}]"
                paragraph = document.add_paragraph()
                run = paragraph.add_run(title)
                run.bold = True
                if code_language == "mermaid":
                    document.add_paragraph("（流程图/时序图，请在 Obsidian 或支持 Mermaid 的工具中查看原 Markdown。）")
                else:
                    for code_line in code_lines:
                        code_paragraph = document.add_paragraph(code_line)
                        code_paragraph.style = "No Spacing"
                code_lines = []
                code_language = ""
            index += 1
            continue

        if in_code_block:
            code_lines.append(line)
            index += 1
            continue

        # 表格
        if stripped.startswith("|"):
            if is_table_separator(stripped):
                index += 1
                continue
            row_cells = parse_table_row(stripped)
            if row_cells:
                table_rows.append(row_cells)
            index += 1
            if index >= len(lines) or not lines[index].strip().startswith("|"):
                add_markdown_table(document, table_rows)
                table_rows = []
            continue

        if table_rows:
            add_markdown_table(document, table_rows)
            table_rows = []

        # 图片
        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if image_match:
            alt_text, image_rel_path = image_match.group(1), image_match.group(2)
            image_abs_path = resolve_image_path(md_dir, image_rel_path)
            if image_abs_path:
                try:
                    document.add_picture(str(image_abs_path), width=Inches(6.0))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    document.add_paragraph(f"[图片加载失败: {image_rel_path}]")
            else:
                document.add_paragraph(f"[图片不存在: {image_rel_path}]")
            if alt_text:
                caption = document.add_paragraph(alt_text)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            index += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2)
            document.add_heading(title_text, level=min(level, 4))
            index += 1
            continue

        # 引用块
        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            paragraph = document.add_paragraph()
            add_rich_text(paragraph, quote_text)
            paragraph.paragraph_format.left_indent = Inches(0.3)
            index += 1
            continue

        # 分隔线
        if stripped in ("---", "***", "___"):
            document.add_paragraph("—" * 40)
            index += 1
            continue

        # 无序列表
        list_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if list_match:
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, list_match.group(1))
            index += 1
            continue

        # 空行
        if not stripped:
            index += 1
            continue

        # 普通段落
        paragraph = document.add_paragraph()
        add_rich_text(paragraph, stripped)
        index += 1

    if table_rows:
        add_markdown_table(document, table_rows)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))


def main() -> None:
    if len(sys.argv) < 3:
        print("用法: py md_to_docx.py <输入.md> <输出.docx>")
        sys.exit(1)
    md_path = Path(sys.argv[1]).resolve()
    docx_path = Path(sys.argv[2]).resolve()
    convert_markdown_to_docx(md_path, docx_path)
    print(f"已生成: {docx_path}")


if __name__ == "__main__":
    main()
